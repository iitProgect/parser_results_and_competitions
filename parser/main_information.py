from yargy import Parser, rule, and_, or_
from yargy.predicates import gram, is_capitalized, dictionary, is_upper, length_eq, is_title
from docx import Document
from docx.shared import Inches
import os, subprocess

baseDir = r'C:\Users\Катя\Desktop\рпд'

def GetDocuments():
    documents = []
    for filename in os.listdir(baseDir):
        if filename.endswith('.docx'):
            documents.append(filename)
    return documents

documents = GetDocuments()
path = baseDir + '\\'+documents[0]
print(path)
document = Document(path)

from yargy import Parser, rule, and_
from yargy.predicates import gram, is_capitalized, dictionary


Name = rule(

    and_(
        gram('ADJF'),
        or_(
        is_capitalized(),
        is_upper()
        )                 # http://pymorphy2.readthedocs.io/en/latest/user/grammemes.html
    ),
    gram('ADJF').optional().repeatable(),
    dictionary({
        'институт',
        'кафедра',
        'университет'
    })
)

Name1 = rule(
    dictionary({
        'институт',
        'кафедра',
        'университет'
    }),
    gram('ADJF'),
    gram('ADJF').optional().repeatable(),
    gram('NOUN'),
)
parser = Parser(Name)
# parser = Parser(Name1)

for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            for match in parser.findall(cell.text):
                print([_.value for _ in match.tokens])



for i in document.paragraphs:
    for match in parser.findall(i.text):
        print([_.value for _ in match.tokens])
