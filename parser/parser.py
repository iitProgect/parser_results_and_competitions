from yargy import Parser, rule, and_, or_
from yargy.predicates import gram, is_capitalized, dictionary, is_upper, length_eq
from docx import Document
from docx.shared import Inches

class CompetitionResult():
    def __init__(self, FGOS, competition, result):
        self.FGOS = FGOS
        self.competition = competition
        self.result = result

FGOS = rule(
     and_(
         is_upper(),
         or_(
            length_eq(2),
            length_eq(3)
         )
        )
)

IsCodeFGOS = rule(
    dictionary(
    {'код', 'ФГОС'})
)
IsCompetitions = rule(dictionary(
    {
    'компетенция'
    }
))
IsResults = rule(dictionary(
    {
        'результат'
    }
))
TableIsCompetitionsAndResults = rule(dictionary(
    {
        'результат'
    }
))

parser_Table = Parser(TableIsCompetitionsAndResults)
parser_FGOS = Parser(IsCodeFGOS)
parser_Result = Parser(IsResults)
parser_Competition = Parser(IsCompetitions)
parser_code_FGOS = Parser(FGOS)

# из текста
def FindFGOSs(text):
    list_code_fgos = []
    for match in parser_code_FGOS.findall(text):
        index = 0
        for i in match.tokens:
            code_FGOS = match.tokens[index].value+text[match.tokens[index].span[1]]+text[match.tokens[index].span[1]+1]
            if text[match.tokens[index].span[1]+2].isdigit():
                code_FGOS = code_FGOS + text[match.tokens[index].span[1]+2]
            list_code_fgos.append(code_FGOS)
            index +=1
    return list_code_fgos

#из строчки
def FindFGOS(text):
    code_FGOS = ""
    for match in parser_code_FGOS.findall(text):
        code_FGOS = match.tokens[0].value+text[match.tokens[0].span[1]]+text[match.tokens[0].span[1]+1]
        if text[match.tokens[0].span[1]+2].isdigit():
            code_FGOS = code_FGOS + text[match.tokens[0].span[1]+2]
    return code_FGOS


# document = Document('C://Users//Катя//Desktop//рпд//РПД Государственная итоговая аттестация.docx')
# document = Document('C://Users//Катя//Desktop//рпд//РПД Аналитика информационных систем (09.03.01, 2016, (4.0), Информатика и вычислительная техника(19610)).docx')
# document = Document('C://Users//Катя//Desktop//рпд//РПД _Дискретная математика по БИ_38.03.05.docx')

# document = Document('C://Users//Катя//Desktop//рпд//4_РПД Экономика.docx')
# document = Document('C://Users//Катя//Desktop//рпд//1.История.docx')
document = Document('C://Users//Катя//Desktop//рпд//2.Философия.docx')

def FindTable():
    for table in document.tables:

        data = []

        keys = None


        if table.rows[0].cells[0].text == table.rows[0].cells[1].text:
            index_row = 0


            for i in range(len(table.rows)-1):
                if index_row == 0:
                    index_row += 1
                    continue
                row_data = {}
                index_column = 0

                for j in table.columns:
                    if index_column == 2 and table.rows[1].cells[index_column].text == '':
                        row_data[table.rows[0].cells[index_column].text] = table.rows[index_row + 1].cells[index_column].text
                    row_data[table.rows[1].cells[index_column].text] = table.rows[index_row+1].cells[index_column].text
                    index_column += 1
                data.append(row_data)
                index_row += 1
            print(data)

        else:
            for i, row in enumerate(table.rows):

                text = (cell.text for cell in row.cells)
                if i == 0:
                    keys = tuple(text)
                    continue
                row_data = dict(zip(keys, text))
                data.append(row_data)

        if len(data) != 0:
            for j in data[0]:
                if parser_Table.find(j) is not None:
                    print(data)
                    return data
        else:
            continue

    return None

FGOS = []
competitions = []
results = []

table = FindTable()

if table is not None:
    #определение какой столбец за что отвечает
    key = 0
    dict={}
    for i in table[0]:
        if parser_FGOS.find(i) is not None and 'ФГОС' not in dict:
            print('ФГОС '+i)
            dict['ФГОС'] = i
        elif parser_Competition.find(i) is not None and 'Компетенция' not in dict:
            print('Компетенция ' + i)
            dict['компетенции'] = i
        elif parser_Result.find(i) is not None and 'Результаты' not in dict:
            print('Результаты ' + i)
            dict['результаты'] = i
    print(dict)
    for i in table:#i - словарь
        if 'компетенции' in dict and not i[dict['компетенции']] == dict['компетенции']:
            competitions.append(i[dict['компетенции']])
        if 'результаты' in dict:
            results.append(i[dict['результаты']])
        if 'ФГОС' in dict:
            FGOS.append(i[dict['ФГОС']])

    print(FGOS)
    print(competitions)
    print(results)
    print('Сущности')
    if len(results) != 0:
        if len(competitions)!= 0 and len(FGOS) != 0:
            competition_result = []
            index_list = 0
            for i in competitions:
                competition_result.append(CompetitionResult(FGOS[index_list],competitions[index_list], results[index_list]))
                index_list += 1
            for i in competition_result:
                print('Код: ' + i.FGOS, 'Компетенция: ' + i.competition, 'Результат: ' + i.result)
        elif len(competitions) == 0 and len(FGOS) != 0:
            print('если коды и компетенции в столбике коды')

        elif len(competitions) != 0 and len(FGOS) == 0:
            print('если коды и компетенции в столбике компетенции')

            competition_result = []
            index_list = 0
            for i in competitions:
                code = FindFGOS(i)
                index = i.find(code)
                competition_result.append(CompetitionResult(code, i[index+len(code):], results[index_list]))
                index_list += 1

            for i in competition_result:
                print('Код: ' + i.FGOS, 'Компетенция: '+ i.competition, 'Результат: '+i.result)




