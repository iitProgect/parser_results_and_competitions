from docx import Document
from docx.shared import Inches
import re
import pymorphy2
from yargy import Parser, rule, and_
from yargy.predicates import gram, is_capitalized, dictionary
morph = pymorphy2.MorphAnalyzer()


def ExistWordFromDictInParagraph(list_words, paragraph):
    words = re.findall(r"[-\w']+", paragraph)
    words = [morph.parse(i)[0].normal_form for i in words]
    return any(elem in list_words for elem in words)

document = Document('C://Users//Катя//Desktop//11_РПД _Программирование (2).docx')



text = document.paragraphs

def FindTable():
    for table in document.tables:

        data = []

        keys = None
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)

            # Establish the mapping based on the first row
            # headers; these will become the keys of our dictionary
            if i == 0:
                keys = tuple(text)
                continue

            # Construct a dictionary for this row, mapping
            # keys to values for this row
            row_data = dict(zip(keys, text))
            data.append(row_data)


            for j in data[0]:
                if ExistWordFromDictInParagraph(['компетенция','результаты','ФГОС'],j):
                    return data

FGOS = []
competitions = []
results = []
table = FindTable()

#определение какой столбец за что отвечает
key = 0
dict={}
for i in table[0]:
    if ExistWordFromDictInParagraph(['код','ФГОС'],i):
        dict['ФГОС'] = 0
    elif ExistWordFromDictInParagraph(['компетенция'],i):
        dict['компетенции'] = 1
    elif ExistWordFromDictInParagraph(['результат'],i):
        dict['результаты'] = 2
    else:
        print(i)
print(dict)






