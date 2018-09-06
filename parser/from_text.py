from docx import Document
from docx.shared import Inches
import re
import pymorphy2
#from yargy import Parser, rule, and_
#from yargy.predicates import gram, is_capitalized, dictionary
morph = pymorphy2.MorphAnalyzer()


def ExistWordFromDictInParagraph(list_words, paragraph):
    words = re.findall(r"[-\w']+", paragraph)
    words = [morph.parse(i)[0].normal_form for i in words]
    return any(elem in list_words for elem in words)

# document = Document('C://Users//Катя//Desktop//11_РПД _Программирование (2).docx')
document = Document('C://Users//Катя//Desktop//ПИГА 09.03.01, 2016, (4.0), Информатика и вычислительная техника (19610).docx')

text = document.paragraphs
index = 0
for i in text:
    print(i.text)
    print(index)
    index+=1
print(text[27].text)
print(text[27].style)
print(text[28].text)
print(text[28].style)
print(text[29].text)
print(text[29].style)


 # def ParagraphIsHeader(index_paragraph):
 #     if
