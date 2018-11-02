from yargy import Parser, rule, and_, or_
from yargy.predicates import gram, is_capitalized, dictionary, is_upper, length_eq, is_title
from docx import Document
from docx.shared import Inches
import os

def GetDocuments():
    documents = []
    for filename in os.listdir(baseDir):
        if filename.endswith('.docx'):
            documents.append(filename)
    return documents

from yargy.tokenizer import TokenRule
from yargy.tokenizer import Tokenizer

CODE_RULE = TokenRule('Code', '\d{2}.\d{2}.\d{2}(?!\d)')
tokenizer = Tokenizer()
tokenizer.remove_types('EOL','LATIN','RU','INT','PUNCT','OTHER')
tokenizer.add_rules(CODE_RULE)

isRPD = rule(
    and_(dictionary({
                'рабочая'}), is_title()),
        dictionary({
                'программа'}))

isRPD2 = rule(
    dictionary({
        'дисциплина'
    })
)

rpdRule = Parser(isRPD)

# print(rpdRule.find('Рабочая программа дисциплины'))


baseDir = r'C:\Users\Katia\Desktop\рпд'
documents = GetDocuments()
path = baseDir + '\\'+documents[0]
print(path)
document = Document(path)
def FindName():
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if len(list(tokenizer(cell.text))) > 0:
                    span = list(tokenizer(cell.text))[0].span
                    return cell.text[span[0]:cell.text.find('\n', span[0])]
    for i in document.paragraphs:
        if len(list(tokenizer(i.text))) == 1:
            return i.text


def FindRPD():
    for row in document.tables[0].rows:
        for cell in row.cells:
            if rpdRule.find(cell.text) is not None:
                return FindRPDInTable(cell.text)
    index = 0
    for i in document.paragraphs:
        if rpdRule.find(i.text) is not None:
            return FindRPDInParagraphs(index+1)
        index += 1


def FindRPDInTable(cell):
    rpdRule = Parser(isRPD2)
    span = rpdRule.find(cell).span
    return cell[span[1]:cell.find('\n', span[1])]

def FindRPDInParagraphs(start_index):
    for i in document.paragraphs:
        if document.paragraphs[start_index].text != '':
            return document.paragraphs[start_index].text
        start_index += 1

print('направление подготовки:',FindName())
print('рабочая программа дисциплины', FindRPD())
